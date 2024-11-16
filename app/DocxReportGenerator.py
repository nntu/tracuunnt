from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from pathlib import Path
from datetime import datetime
import logging
from typing import Dict, List, Union, Optional
from PIL import Image

class DocxReportGenerator:
    """Generates Word document reports with screenshots"""
    
    def __init__(self, save_dir: Union[str, Path]):
        self.save_dir = Path(save_dir)
        self.save_dir.mkdir(parents=True, exist_ok=True)
        
    def create_docx_report(self, 
                          result_df: pd.DataFrame,
                          title: str = "Invoice Check Report",
                          screenshots: Optional[Dict[str, str]] = None) -> Path:
        """
        Creates a Word document with results and screenshots
        """
        try:
            # Create new document
            doc = Document()
            
            # Set landscape orientation and margins
            section = doc.sections[0]
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            
            # Set small margins for maximum space
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add datetime
            date_paragraph = doc.add_paragraph()
            date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_run.font.size = Pt(10)
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary table
            if not result_df.empty:
                summary_table = doc.add_table(rows=1, cols=2)
                summary_table.style = 'Table Grid'
                header_cells = summary_table.rows[0].cells
                header_cells[0].text = 'Metric'
                header_cells[1].text = 'Value'
                
                summary_data = {
                    'Total Records': len(result_df),
                    'Total Screenshots': len(screenshots) if screenshots else 0,
                    'Processing Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
                
                for key, value in summary_data.items():
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = key
                    row_cells[1].text = str(value)
            
            doc.add_page_break()
            
            # Process each MST if screenshots are provided
            if screenshots:
                for mst, screenshot_path in screenshots.items():
                    screenshot_path = Path(screenshot_path)
                    if screenshot_path.exists():
                        # Create main table for the page
                        main_table = doc.add_table(rows=3, cols=1)
                        main_table.style = 'Table Grid'
                        
                        # MST Header cell
                        header_cell = main_table.rows[0].cells[0]
                        header_paragraph = header_cell.paragraphs[0]
                        header_run = header_paragraph.add_run(f"MST: {mst}")
                        header_run.font.bold = True
                        header_run.font.size = Pt(12)
                        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Information cell
                        info_cell = main_table.rows[1].cells[0]
                        if not result_df.empty and 'MST' in result_df.columns:
                            mst_info = result_df[result_df['MST'] == mst]
                            if not mst_info.empty:
                                info_table = info_cell.add_table(rows=len(mst_info.columns)-1, cols=2)
                                info_table.style = 'Table Grid'
                                
                                row_idx = 0
                                for col in mst_info.columns:
                                    if col != 'MST':
                                        cells = info_table.rows[row_idx].cells
                                        cells[0].text = str(col)
                                        cells[1].text = str(mst_info[col].iloc[0])
                                        row_idx += 1
                        
                        # Screenshot cell
                        img_cell = main_table.rows[2].cells[0]
                        
                        # Calculate image dimensions
                        with Image.open(screenshot_path) as img:
                            width, height = img.size
                            
                        # Calculate scaling to fit page
                        max_width = Inches(10)    # Landscape page width minus margins
                        max_height = Inches(4.5)  # Leave space for MST and info
                        
                        width_scale = max_width / width
                        height_scale = max_height / height
                        scale = min(width_scale, height_scale)
                        
                        new_width = int(width * scale)
                        new_height = int(height * scale)
                        
                        # Add image
                        img_paragraph = img_cell.paragraphs[0]
                        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_run = img_paragraph.add_run()
                        img_run.add_picture(
                            str(screenshot_path),
                            width=new_width,
                            height=new_height
                        )
                        
                        # Set row heights
                        main_table.rows[0].height = Inches(0.4)  # MST header
                        main_table.rows[1].height = Inches(2.0)  # Info
                        main_table.rows[2].height = Inches(5.0)  # Screenshot
                        
                        # Add page break after each entry except the last one
                        if mst != list(screenshots.keys())[-1]:
                            doc.add_page_break()
            
            # Save documents
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            docx_path = self.save_dir / f"report_{timestamp}.docx"
            excel_path = self.save_dir / f"report_{timestamp}.xlsx"
            
            doc.save(str(docx_path))
            result_df.to_excel(str(excel_path), index=False)
            
            logging.info(f"Created reports at: {self.save_dir}")
            
            return docx_path
            
        except Exception as e:
            logging.error(f"Failed to create reports: {str(e)}")
            raise